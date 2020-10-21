# Copyright (C) 2019 by Sukchan Lee <acetcom@gmail.com>

# This file is part of Open5GS.

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

from docx import Document
import re, os, sys, string
import datetime
import getopt
import getpass

version = "0.1.0"

msg_list = {}
type_list = {}
group_list = {}

verbosity = 0
filename = ""
outdir = './'
cachedir = './cache/'

FAIL = '\033[91m'
INFO = '\033[93m'
ENDC = '\033[0m'

def d_print(string):
    if verbosity > 0:
        sys.stdout.write(string)

def d_info(string):
    sys.stdout.write(INFO + string + ENDC + "\n")

def d_error(string):
    sys.stderr.write(FAIL + string + ENDC + "\n")
    sys.exit(0)

def write_file(f, string):
    f.write(string)
    d_print(string)

def output_header_to_file(f):
    now = datetime.datetime.now()
    f.write("""/*
 * Copyright (C) 2019 by Sukchan Lee <acetcom@gmail.com>
 *
 * This file is part of Open5GS.
 *
 * This program is free software: you can redistribute it and/or modify
 * it under the terms of the GNU Affero General Public License as published by
 * the Free Software Foundation, either version 3 of the License, or
 * (at your option) any later version.
 *
 * This program is distributed in the hope that it will be useful,
 * but WITHOUT ANY WARRANTY; without even the implied warranty of
 * MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
 * GNU General Public License for more details.
 *
 * You should have received a copy of the GNU General Public License
 * along with this program.  If not, see <https://www.gnu.org/licenses/>.
 */

""")
    f.write("/*******************************************************************************\n")
    f.write(" * This file had been created by gtp-tlv.py script v%s\n" % (version))
    f.write(" * Please do not modify this file but regenerate it via script.\n")
    f.write(" * Created on: %s by %s\n * from %s\n" % (str(now), getpass.getuser(), filename))
    f.write(" ******************************************************************************/\n\n")

def usage():
    print "Python generating TLV build/parser for GTPv2-C v%s" % (version)
    print "Usage: python gtp-tlv.py [options]"
    print "Available options:"
    print "-d        Enable script debug"
    print "-f [file] Input file to parse"
    print "-o [dir]  Output files to given directory"
    print "-c [dir]  Cache files to given directory"
    print "-h        Print this help and return"

def v_upper(v):
    return re.sub('3GPP', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).upper())

def v_lower(v):
    return re.sub('3gpp', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).lower())

def get_cells(cells):
    instance = cells[4].text.encode('ascii', 'ignore')
    if instance.isdigit() is not True:
        return None
    ie_type = re.sub('\s*$', '', re.sub('\s*\n*\s*\([A-z]*\s*NOTE.*\)*', '', cells[3].text.encode('ascii', 'ignore')))
    if ie_type.find('LDN') != -1:
        ie_type = 'LDN'
    elif ie_type.find('APCO') != -1:
        ie_type = 'APCO'
    elif ie_type.find('Charging Id') != -1:
        ie_type = 'Charging ID'
    elif ie_type.find('H(e)NB Information Reporting') != -1:
        ie_type = 'eNB Information Reporting'
    elif ie_type.find('IPv4 Configuration Parameters (IP4CP)') != -1:
        ie_type = 'IP4CP'
    elif ie_type.find('Charging characteristics') != -1:
        ie_type = 'Charging Characteristics'
    elif ie_type.find('Change To Report Flags') != -1:
        ie_type = 'Change to Report Flags'
    elif ie_type.find('APN RATE Control Status') != -1:
        ie_type = 'APN Rate Control Status'
    if ie_type not in type_list.keys():
        assert False, "Unknown IE type : [" \
                + cells[3].text + "]" + "(" + ie_type + ")"
    presence = cells[1].text.encode('ascii', 'ignore')
    presence = re.sub('\n', '', presence);
    ie_value = re.sub('\s*\n*\s*\([^\)]*\)*', '', cells[0].text).encode('ascii', 'ignore')
    comment = cells[2].text.encode('ascii', 'ignore')
    comment = re.sub('\n|\"|\'|\\\\', '', comment);

    if int(instance) > int(type_list[ie_type]["max_instance"]):
        type_list[ie_type]["max_instance"] = instance
        write_file(f, "type_list[\"" + ie_type + "\"][\"max_instance\"] = \"" + instance + "\"\n")

    return { "ie_type" : ie_type, "ie_value" : ie_value, "presence" : presence, "instance" : instance, "comment" : comment }

def write_cells_to_file(name, cells):
    write_file(f, name + ".append({ \"ie_type\" : \"" + cells["ie_type"] + \
        "\", \"ie_value\" : \"" + cells["ie_value"] + \
        "\", \"presence\" : \"" + cells["presence"] + \
        "\", \"instance\" : \"" + cells["instance"] + \
        "\", \"comment\" : \"" + cells["comment"] + "\"})\n")

try:
    opts, args = getopt.getopt(sys.argv[1:], "df:ho:c:", ["debug", "file", "help", "output", "cache"])
except getopt.GetoptError as err:
    # print help information and exit:
    usage()
    sys.exit(2)

for o, a in opts:
    if o in ("-d", "--debug"):
        verbosity = 1
    if o in ("-f", "--file"):
        filename = a
    if o in ("-o", "--output"):
        outdir = a
        if outdir.rfind('/') != len(outdir):
            outdir += '/'
    if o in ("-c", "--cache"):
        cache = a
        if cachedir.rfind('/') != len(cachedir):
            cachedir += '/'
    if o in ("-h", "--help"):
        usage()
        sys.exit(2)

if os.path.isfile(filename) and os.access(filename, os.R_OK):
    file = open(filename, 'r') 
else:
    d_error("Cannot find file : " + filename)

d_info("[Message List]")
cachefile = cachedir + 'tlv-msg-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    msg_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('Message Type value') != -1:
            msg_table = table
            d_print("Table Index = %d\n" % i)

    for row in msg_table.rows[2:-4]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        type = row.cells[0].text.encode('ascii', 'ignore')
        if type.isdigit() is False:
            continue
        if int(type) in range(128, 160):
            continue
        if int(type) in range(231, 240):
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\s*\n*\s*\([^\)]*\)*', '', key)
        msg_list[key] = { "type": type }
        write_file(f, "msg_list[\"" + key + "\"] = { \"type\" : \"" + type + "\" }\n")
    f.close()

d_info("[IE Type List]")
cachefile = cachedir + 'tlv-type-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    ie_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('IE Type value') != -1:
            ie_table = table
            d_print("Table Index = %d\n" % i)

    for row in ie_table.rows[1:-5]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        if key.find('Reserved') != -1:
            continue
        if key.find('MM Context') != -1:
            continue
        elif key.find('Recovery') != -1:
            key = 'Recovery'
        elif key.find('Trusted WLAN Mode Indication') != -1:
            key = 'TWMI'
        elif key.find('LDN') != -1:
            key = 'LDN'
        elif key.find('APCO') != -1:
            key = 'APCO'
        elif key.find('Remote UE IP information') != -1:
            key = 'Remote UE IP Information'
        elif key.find('Procedure Transaction ID') != -1:
            key = 'PTI'
        else:
            key = re.sub('.*\(', '', row.cells[1].text.encode('ascii', 'ignore'))
            key = re.sub('\)', '', key)
            key = re.sub('\s*$', '', key)
        type = row.cells[0].text.encode('ascii', 'ignore')
        type_list[key] = { "type": type , "max_instance" : "0" }
        write_file(f, "type_list[\"" + key + "\"] = { \"type\" : \"" + type)
        write_file(f, "\", \"max_instance\" : \"0\" }\n")
    f.close()
type_list['MM Context'] = { "type": "107", "max_instance" : "0" }

d_info("[Group IE List]")
cachefile = cachedir + 'tlv-group-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    for i, table in enumerate(document.tables):
        if table.rows[0].cells[0].text.find('Octet') != -1 and \
            table.rows[0].cells[2].text.find('IE Type') != -1:
            d_print("Table Index = %d\n" % i)

            row = table.rows[0];

            if len(re.findall('\d+', row.cells[2].text)) == 0:
                continue;
            ie_type = re.findall('\d+', row.cells[2].text)[0].encode('ascii', 'ignore')
            ie_name = re.sub('\s*IE Type.*', '', row.cells[2].text.encode('ascii', 'ignore'))

            if ie_name not in group_list.keys():
                ies = []
                write_file(f, "ies = []\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    ies_is_added = True
                    for ie in ies:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        ies.append(cells)
                        write_cells_to_file("ies", cells)

                ie_idx = str(int(ie_type)+100)
                group_list[ie_name] = { "index" : ie_idx, "type" : ie_type, "ies" : ies }
                write_file(f, "group_list[\"" + ie_name + "\"] = { \"index\" : \"" + ie_idx + "\", \"type\" : \"" + ie_type + "\", \"ies\" : ies }\n")
            else:
                group_list_is_added = False
                added_ies = group_list[ie_name]["ies"]
                write_file(f, "added_ies = group_list[\"" + ie_name + "\"][\"ies\"]\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    ies_is_added = True
                    for ie in group_list[ie_name]["ies"]:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    for ie in ies:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        added_ies.append(cells)
                        write_cells_to_file("added_ies", cells)
                        group_list_is_added = True
                if group_list_is_added is True:
                    ie_idx = str(int(ie_type)+100)
                    group_list[ie_name] = { "index" : ie_idx, "type" : ie_type, "ies" : added_ies }
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"index\" : \"" + ie_idx + "\", \"type\" : \"" + ie_type + "\", \"ies\" : added_ies }\n")
    f.close()

msg_list["Echo Request"]["table"] = 8
msg_list["Echo Response"]["table"] = 9
msg_list["Create Session Request"]["table"] = 10
msg_list["Create Session Response"]["table"] = 15
msg_list["Create Bearer Request"]["table"] = 20
msg_list["Create Bearer Response"]["table"] = 24
msg_list["Bearer Resource Command"]["table"] = 27
msg_list["Bearer Resource Failure Indication"]["table"] = 29
msg_list["Modify Bearer Request"]["table"] = 31
msg_list["Modify Bearer Response"]["table"] = 35
msg_list["Delete Session Request"]["table"] = 40
msg_list["Delete Bearer Request"]["table"] = 42
msg_list["Delete Session Response"]["table"] = 46
msg_list["Delete Bearer Response"]["table"] = 49
msg_list["Downlink Data Notification"]["table"] = 52
msg_list["Downlink Data Notification Acknowledge"]["table"] = 55
msg_list["Downlink Data Notification Failure Indication"]["table"] = 56
msg_list["Delete Indirect Data Forwarding Tunnel Request"]["table"] = 57
msg_list["Delete Indirect Data Forwarding Tunnel Response"]["table"] = 58
msg_list["Modify Bearer Command"]["table"] = 59
msg_list["Modify Bearer Failure Indication"]["table"] = 62
msg_list["Update Bearer Request"]["table"] = 64
msg_list["Update Bearer Response"]["table"] = 68
msg_list["Delete Bearer Command"]["table"] = 71
msg_list["Delete Bearer Failure Indication"]["table"] = 74
msg_list["Create Indirect Data Forwarding Tunnel Request"]["table"] = 77
msg_list["Create Indirect Data Forwarding Tunnel Response"]["table"] = 79
msg_list["Release Access Bearers Request"]["table"] = 81
msg_list["Release Access Bearers Response"]["table"] = 82
msg_list["Modify Access Bearers Request"]["table"] = 86
msg_list["Modify Access Bearers Response"]["table"] = 89

for key in msg_list.keys():
    if "table" in msg_list[key].keys():
        d_info("[" + key + "]")
        cachefile = cachedir + "tlv-msg-" + msg_list[key]["type"] + ".py"
        if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
            execfile(cachefile)
            print "Read from " + cachefile
        else:
            document = Document(filename)
            f = open(cachefile, 'w') 

            ies = []
            write_file(f, "ies = []\n")
            table = document.tables[msg_list[key]["table"]]
            for row in table.rows[1:]:
                cells = get_cells(row.cells)
                if cells is None:
                    continue

                ies_is_added = True
                for ie in ies:
                    if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                        ies_is_added = False
                if ies_is_added is True:
                    ies.append(cells)
                    write_cells_to_file("ies", cells)
            msg_list[key]["ies"] = ies
            write_file(f, "msg_list[key][\"ies\"] = ies\n")
            f.close()

type_list["Recovery"]["size"] = 1                       # Type : 3
type_list["EBI"]["size"] = 1                            # Type : 73
type_list["RAT Type"]["size"] = 1                       # Type : 82
type_list["PDN Type"]["size"] = 1                       # Type : 99
type_list["PTI"]["size"] = 1                            # Type : 100
type_list["Port Number"]["size"] = 2                    # Type : 126
type_list["APN Restriction"]["size"] = 1                # Type : 127
type_list["Selection Mode"]["size"] = 1                 # Type : 128
type_list["Node Type"]["size"] = 1                      # Type : 135
type_list["Node Features"]["size"] = 1                  # Type : 152

f = open(outdir + 'message.h', 'w')
output_header_to_file(f)
f.write("""#if !defined(OGS_GTP_INSIDE) && !defined(OGS_GTP_COMPILATION)
#error "This header cannot be included directly."
#endif

#ifndef OGS_GTP_MESSAGE_H
#define OGS_GTP_MESSAGE_H

#ifdef __cplusplus
extern "C" {
#endif

/* 5.1 General format */
#define OGS_GTPV1U_HEADER_LEN   8
#define OGS_GTPV2C_HEADER_LEN   12
#define OGS_GTP_TEID_LEN        4
typedef struct ogs_gtp_header_s {
    union {
        struct {
#define OGS_GTP_VERSION_0 0
#define OGS_GTP_VERSION_1 1
        ED4(uint8_t version:3;,
            uint8_t piggybacked:1;,
            uint8_t teid_presence:1;,
            uint8_t spare1:3;)
        };
/* GTU-U flags */
#define OGS_GTPU_FLAGS_V                        0x20
#define OGS_GTPU_FLAGS_PT                       0x10
#define OGS_GTPU_FLAGS_E                        0x04
#define OGS_GTPU_FLAGS_S                        0x02
#define OGS_GTPU_FLAGS_PN                       0x01
        uint8_t flags;
    };
/* GTP-U message type, defined in 3GPP TS 29.281 Release 11 */
#define OGS_GTPU_MSGTYPE_ECHO_REQ               1
#define OGS_GTPU_MSGTYPE_ECHO_RSP               2
#define OGS_GTPU_MSGTYPE_ERR_IND                26
#define OGS_GTPU_MSGTYPE_SUPP_EXTHDR_NOTI       31
#define OGS_GTPU_MSGTYPE_END_MARKER             254
#define OGS_GTPU_MSGTYPE_GPDU                   255
    uint8_t type;
    uint16_t length;
    union {
        struct {
            uint32_t teid;
            /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
#define OGS_GTP_XID_TO_SQN(__xid) htobe32(((__xid) << 8))
#define OGS_GTP_SQN_TO_XID(__sqn) (be32toh(__sqn) >> 8)
            uint32_t sqn;
        };
        /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
        uint32_t sqn_only;
    };
} __attribute__ ((packed)) ogs_gtp_header_t;

/* GTPv2-C message type */
""")

tmp = [(k, v["type"]) for k, v in msg_list.items()]
sorted_msg_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_msg_list:
    f.write("#define OGS_GTP_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in type_list.items()]
sorted_type_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_type_list:
    f.write("#define OGS_GTP_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

f.write("/* Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("extern ogs_tlv_desc_t ogs_gtp_tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

for k, v in group_list.items():
    if v_lower(k) == "pc5_qos_parameters":
        v["index"] = "1"
    if v_lower(k) == "remote_ue_context":
        v["index"] = "2"

tmp = [(k, v["index"]) for k, v in group_list.items()]
sorted_group_list = sorted(tmp, key=lambda tup: int(tup[1]))

f.write("/* Group Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("extern ogs_tlv_desc_t ogs_gtp_tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

f.write("/* Message Descriptor */\n")
for (k, v) in sorted_msg_list:
    f.write("extern ogs_tlv_desc_t ogs_gtp_tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Structure for Infomration Element */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    if "size" in type_list[k]:
        if type_list[k]["size"] == 1:
            f.write("typedef ogs_tlv_uint8_t ogs_gtp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 2:
            f.write("typedef ogs_tlv_uint16_t ogs_gtp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 3:
            f.write("typedef ogs_tlv_uint24_t ogs_gtp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 4:
            f.write("typedef ogs_tlv_uint32_t ogs_gtp_tlv_" + v_lower(k) + "_t;\n")
        else:
            assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
    else:
        f.write("typedef ogs_tlv_octet_t ogs_gtp_tlv_" + v_lower(k) + "_t;\n")
f.write("\n")

f.write("/* Structure for Group Infomration Element */\n")
for (k, v) in sorted_group_list:
    f.write("typedef struct ogs_gtp_tlv_" + v_lower(k) + "_s {\n")
    f.write("    ogs_tlv_presence_t presence;\n")
    for ies in group_list[k]["ies"]:
        f.write("    ogs_gtp_tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                v_lower(ies["ie_value"]))
        if ies["ie_type"] == "F-TEID":
            if ies["ie_value"] == "S2b-U ePDG F-TEID":
                f.write("_" + ies["instance"] + ";")
            elif ies["ie_value"] == "S2a-U TWAN F-TEID":
                f.write("_" + ies["instance"] + ";")
            else:
                f.write(";")
            f.write(" /* Instance : " + ies["instance"] + " */\n")
        else:
            f.write(";\n")
    f.write("} ogs_gtp_tlv_" + v_lower(k) + "_t;\n")
    f.write("\n")

f.write("/* Structure for Message */\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("typedef struct ogs_gtp_" + v_lower(k) + "_s {\n")
        for ies in msg_list[k]["ies"]:
            if (k == 'Create Indirect Data Forwarding Tunnel Request' or k == 'Create Indirect Data Forwarding Tunnel Response') and ies["ie_value"] == 'Bearer Contexts':
                f.write("    ogs_gtp_tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                        v_lower(ies["ie_value"]) + "[8];\n")
            else:
                f.write("    ogs_gtp_tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                        v_lower(ies["ie_value"]) + ";\n")
        f.write("} ogs_gtp_" + v_lower(k) + "_t;\n")
        f.write("\n")

f.write("typedef struct ogs_gtp_message_s {\n")
f.write("   ogs_gtp_header_t h;\n")
f.write("   union {\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        ogs_gtp_" + v_lower(k) + "_t " + v_lower(k) + ";\n");
f.write("   };\n");
f.write("} ogs_gtp_message_t;\n\n")

f.write("""int ogs_gtp_parse_msg(ogs_gtp_message_t *gtp_message, ogs_pkbuf_t *pkbuf);
ogs_pkbuf_t *ogs_gtp_build_msg(ogs_gtp_message_t *gtp_message);

#ifdef __cplusplus
}
#endif

#endif /* OGS_GTP_MESSAGE_H */
""")
f.close()

f = open(outdir + 'message.c', 'w')
output_header_to_file(f)
f.write("""#include "ogs-gtp.h"

""")

for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("ogs_tlv_desc_t ogs_gtp_tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        if "size" in type_list[k]:
            if type_list[k]["size"] == 1:
                f.write("    OGS_TLV_UINT8,\n")
            elif type_list[k]["size"] == 2:
                f.write("    OGS_TLV_UINT16,\n")
            elif type_list[k]["size"] == 3:
                f.write("    OGS_TLV_UINT24,\n")
            elif type_list[k]["size"] == 4:
                f.write("    OGS_TLV_UINT32,\n")
            else:
                assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
        else:
            f.write("    OGS_TLV_VAR_STR,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    OGS_GTP_%s_TYPE,\n" % v_upper(k))
        if "size" in type_list[k]:
            f.write("    %d,\n" % type_list[k]["size"])
        else:
            f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(ogs_gtp_tlv_%s_t),\n" % v_lower(k))
        f.write("    { NULL }\n")
        f.write("};\n\n")

for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("ogs_tlv_desc_t ogs_gtp_tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        f.write("    OGS_TLV_COMPOUND,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    OGS_GTP_%s_TYPE,\n" % v_upper(k))
        f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(ogs_gtp_tlv_%s_t),\n" % v_lower(k))
        f.write("    {\n")
        for ies in group_list[k]["ies"]:
                f.write("        &ogs_gtp_tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
        f.write("        NULL,\n")
        f.write("    }\n")
        f.write("};\n\n")

for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("ogs_tlv_desc_t ogs_gtp_tlv_desc_%s =\n" % v_lower(k))
        f.write("{\n")
        f.write("    OGS_TLV_MESSAGE,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    0, 0, 0, 0, {\n")
        for ies in msg_list[k]["ies"]:
            f.write("        &ogs_gtp_tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
            if (k == 'Create Indirect Data Forwarding Tunnel Request' or k == 'Create Indirect Data Forwarding Tunnel Response') and ies["ie_value"] == 'Bearer Contexts':
                f.write("        &ogs_tlv_desc_more8,\n")
        f.write("    NULL,\n")
        f.write("}};\n\n")
f.write("\n")

f.write("""int ogs_gtp_parse_msg(ogs_gtp_message_t *gtp_message, ogs_pkbuf_t *pkbuf)
{
    int rv = OGS_ERROR;
    ogs_gtp_header_t *h = NULL;
    uint16_t size = 0;

    ogs_assert(gtp_message);
    ogs_assert(pkbuf);
    ogs_assert(pkbuf->len);

    h = (ogs_gtp_header_t *)pkbuf->data;
    ogs_assert(h);
    
    memset(gtp_message, 0, sizeof(ogs_gtp_message_t));

    if (h->teid_presence)
        size = OGS_GTPV2C_HEADER_LEN;
    else
        size = OGS_GTPV2C_HEADER_LEN-OGS_GTP_TEID_LEN;

    ogs_assert(ogs_pkbuf_pull(pkbuf, size));
    memcpy(&gtp_message->h, pkbuf->data - size, size);

    if (h->teid_presence)
        gtp_message->h.teid = be32toh(gtp_message->h.teid);

    if (pkbuf->len == 0) {
        ogs_assert(ogs_pkbuf_push(pkbuf, size));
        return OGS_OK;
    }

    switch(gtp_message->h.type) {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("    case OGS_GTP_%s_TYPE:\n" % v_upper(k))
        f.write("        rv = ogs_tlv_parse_msg(&gtp_message->%s,\n" % v_lower(k))
        f.write("                &ogs_gtp_tlv_desc_%s, pkbuf, OGS_TLV_MODE_T1_L2_I1);\n" % v_lower(k))
        f.write("        break;\n")
f.write("""    default:
        ogs_warn("Not implmeneted(type:%d)", gtp_message->h.type);
        break;
    }

    ogs_assert(ogs_pkbuf_push(pkbuf, size));

    return rv;
}

""")

f.write("""ogs_pkbuf_t *ogs_gtp_build_msg(ogs_gtp_message_t *gtp_message)
{
    ogs_pkbuf_t *pkbuf = NULL;

    ogs_assert(gtp_message);
    switch(gtp_message->h.type) {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("    case OGS_GTP_%s_TYPE:\n" % v_upper(k))
        f.write("        pkbuf = ogs_tlv_build_msg(&ogs_gtp_tlv_desc_%s,\n" % v_lower(k))
        f.write("                &gtp_message->%s, OGS_TLV_MODE_T1_L2_I1);\n" % v_lower(k))
        f.write("        break;\n")
f.write("""    default:
        ogs_warn("Not implmeneted(type:%d)", gtp_message->h.type);
        break;
    }

    return pkbuf;
}
""")

f.write("\n")

f.close()
