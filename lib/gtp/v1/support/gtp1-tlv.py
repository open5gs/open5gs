# Copyright (C) 2019 by Sukchan Lee <acetcom@gmail.com>
# Copyright (C) 2022 by sysmocom - s.f.m.c. GmbH <info@sysmocom.de>

# This file is part of Open5GS.

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

import re, os, sys, string
import datetime
import getopt
import getpass

version = "0.1.0"

msg_list = {}
type_list = {}

verbosity = 0
filename = ""
outdir = './'
cachedir = './cache/'

FAIL = '\033[91m'
INFO = '\033[93m'
ENDC = '\033[0m'

def d_print(string):
    if verbosity > 0:
        sys.stdout.write(string)

def d_info(string):
    sys.stdout.write(INFO + string + ENDC + "\n")

def d_error(string):
    sys.stderr.write(FAIL + string + ENDC + "\n")
    sys.exit(0)

def write_file(f, string):
    f.write(string)
    d_print(string)

def output_header_to_file(f):
    now = datetime.datetime.now()
    f.write("""/*
 * Copyright (C) 2019 by Sukchan Lee <acetcom@gmail.com>
 * Copyright (C) 2022 by sysmocom - s.f.m.c. GmbH <info@sysmocom.de>
 *
 * This file is part of Open5GS.
 *
 * This program is free software: you can redistribute it and/or modify
 * it under the terms of the GNU Affero General Public License as published by
 * the Free Software Foundation, either version 3 of the License, or
 * (at your option) any later version.
 *
 * This program is distributed in the hope that it will be useful,
 * but WITHOUT ANY WARRANTY; without even the implied warranty of
 * MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
 * GNU General Public License for more details.
 *
 * You should have received a copy of the GNU General Public License
 * along with this program.  If not, see <https://www.gnu.org/licenses/>.
 */

""")
    f.write("/*******************************************************************************\n")
    f.write(" * This file had been created by gtp1-tlv.py script v%s\n" % (version))
    f.write(" * Please do not modify this file but regenerate it via script.\n")
    f.write(" * Created on: %s by %s\n * from %s\n" % (str(now), getpass.getuser(), filename))
    f.write(" ******************************************************************************/\n\n")

def usage():
    print("Python generating TLV build/parser for GTPv2-C v%s" % (version))
    print("Usage: python gtp-tlv.py [options]")
    print("Available options:")
    print("-d        Enable script debug")
    print("-f [file] Input file to parse")
    print("-o [dir]  Output files to given directory")
    print("-c [dir]  Cache files to given directory")
    print("-h        Print this help and return")

def v_upper(v):
    return re.sub('3GPP', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).upper())

def v_lower(v):
    return re.sub('3gpp', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).lower())

def ie_reference2type(reference):
    ie_type = 'foobar'
    for k, v in type_list.items():
        if v["reference"] == reference:
            return k
    raise ValueError('no IE type found for reference \"%s\"' % reference)

def get_cells(cells):
    if (len(cells) < 3):
        #"FIXME: Table 7.5A.9: Information Elements in a Delete MBMS Context Request" format in document is broken:
        d_info("Expected length 3: %r" % repr(cells))
        return None
    presence = cells[1].text
    presence = re.sub('\n', '', presence);
    ie_value = re.sub('\s*\n*\s*\([^\)]*\)*', '', cells[0].text)
    ie_value = re.sub('\\xa0', ' ', ie_value) # drop unicode char "No-Break Space" in "Higher bitrates than 16 Mbps flag"
    comment = cells[2].text.encode('ascii', 'ignore').decode('utf-8').rstrip()
    comment = re.sub('\n|\"|\'|\\\\', '', comment);
    if comment == 'GSN Address 7.7.32':
        reference = '7.7.32'
    else:
        reference = comment

    if ie_value == '' and reference == '7.7.16':
        # For some unknown reason "cells[0].text" is '' in this row
        ie_value = 'Teardown Ind'

    return { "ie_value" : ie_value, "presence" : presence, "reference": reference, }

def write_cells_to_file(name, cells):
    write_file(f, name + ".append({ " + \
        "\"ie_value\" : \"" + cells["ie_value"] + \
        "\", \"presence\" : \"" + cells["presence"] + \
        "\", \"reference\" : \"" + cells["reference"] + "\"})\n")

def document_paragraph_tables(document):

    tables = []
    # iterate .docx objects
    def iter_block_items(parent):

      if isinstance(parent, _Document):
          parent_elm = parent.element.body
      elif isinstance(parent, _Cell):
          parent_elm = parent._tc
      elif isinstance(parent, _Row):
          parent_elm = parent._tr
      else:
          raise ValueError("Document format error.")

      for child in parent_elm.iterchildren():
          if isinstance(child, CT_P):
              yield Paragraph(child, parent)
          elif isinstance(child, CT_Tbl):
              yield Table(child, parent)

    idx = -1
    paragraph = ''
    for block in iter_block_items(document):
        table=[]
        # memorize the paragraph
        if isinstance(block, Paragraph):
          paragraph = block.text
          continue
        # fetch the table
        if isinstance(block, Table):
            idx += 1
            table = block
        # store table having a paragraph name
        tables.append([idx, paragraph, table])

    return tables


try:
    opts, args = getopt.getopt(sys.argv[1:], "df:ho:c:", ["debug", "file", "help", "output", "cache"])
except getopt.GetoptError as err:
    # print help information and exit:
    usage()
    sys.exit(2)

for o, a in opts:
    if o in ("-d", "--debug"):
        verbosity = 1
    if o in ("-f", "--file"):
        filename = a
    if o in ("-o", "--output"):
        outdir = a
        if outdir.rfind('/') != len(outdir):
            outdir += '/'
    if o in ("-c", "--cache"):
        cachedir = a
        if cachedir.rfind('/') != len(cachedir):
            cachedir += '/'
    if o in ("-h", "--help"):
        usage()
        sys.exit(2)

if os.path.isfile(filename) and os.access(filename, os.R_OK):
    file = open(filename, 'r')
else:
    d_error("Cannot find file : " + filename)

d_info("[Message List]")
cachefile = cachedir + 'tlv-msg-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    exec(open(cachefile).read())
    print("Read from " + cachefile)
else:
    document = Document(filename)
    f = open(cachefile, 'w')

    msg_table = ""
    for i, paragraph, table in document_paragraph_tables(document):
        cell = table.rows[0].cells[0]
        if cell.text.find('Message Type value') != -1:
            msg_table = table
            d_print("Table Index = %d Name = [%s]\n" % (i, paragraph))
            write_file(f, "# [%s] Index = %d\n" % (paragraph, i))

    for row in msg_table.rows[2:-6]:
        key = row.cells[1].text
        type = row.cells[0].text
        if type.isdigit() is False:
            continue
        if int(type) in range(128, 160):
            continue
        if int(type) in range(231, 240):
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\s*\n*\s*\([^\)]*\)*', '', key)
        msg_list[key] = { "type": type }
        write_file(f, "msg_list[\"" + key + "\"] = { \"type\" : \"" + type + "\" }\n")
    f.close()

d_info("[IE Type List]")
cachefile = cachedir + 'tlv-type-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    exec(open(cachefile).read())
    print("Read from " + cachefile)
else:
    document = Document(filename)
    f = open(cachefile, 'w')

    ie_table = ""
    for i, paragraph, table in document_paragraph_tables(document):
        d_print("iter: Table Index = %d Name = [%s]\n" % (i, paragraph))
        cell = table.rows[0].cells[0]
        if cell.text.find('IE Type Value') != -1:
            ie_table = table
            d_print("Table Index = %d Name = [%s]\n" % (i, paragraph))
            write_file(f, "# [%s] Index = %d\n" % (paragraph, i))

    for row in ie_table.rows[1:-3]:
        type = row.cells[0].text
        format = row.cells[1].text
        key = row.cells[2].text
        reference = row.cells[3].text
        len_type = row.cells[4].text
        if key.find('Reserved') != -1:
            continue
        if key.find('Spare') != -1:
            continue
        else:
            key = re.sub('.*\(', '', row.cells[2].text)
            key = re.sub('\)', '', key)
            key = re.sub('\s*$', '', key)
            if key == '' and type == '19':
                # For some unknown reason "row.cells[2].text" is '' in this row
                key = 'Teardown Ind'
        type_list[key] = { 'type': type , 'reference': reference, 'format': format }
        write_file(f, "type_list[\"" + key + "\"] = { \"type\" : \"" + type)
        write_file(f, "\", \"reference\" : \"" + reference)
        write_file(f, "\", \"format\" : \"" + format)
        if (format.find('TLV') != -1 or format.find('TV') != -1) and len_type.find('Fixed') != -1:
            size = int(row.cells[5].text)
            type_list[key]['size'] = size
            write_file(f, "\", \"size\" : " + str(size))
        else:
            write_file(f, "\"")
        write_file(f, " }\n")
    f.close()
    #type_list['MM Context'] = { "type": "107", "reference": "7.7.28" }


def set_c_type(ie_name, c_type=''):
    if len(c_type):
        type_list[ie_name]["c_type"] = c_type

def set_size(ie_name, size, c_type=''):
    type_list[ie_name]["size"] = size

set_size("MS Time Zone", 2) # Wrongly specified as 1 in spec table
set_c_type("Cause", 'uint')
set_c_type("TLLI", 'uint')
set_c_type("P-TMSI", 'uint')
set_c_type("Recovery", 'uint')
set_c_type("Tunnel Endpoint Identifier Data I", 'uint')
set_c_type("Tunnel Endpoint Identifier Control Plane", 'uint')
set_c_type("Teardown Ind", 'uint')
set_c_type("NSAPI", 'uint')
set_c_type("RAT Type", 'uint')
set_c_type("APN Restriction", 'uint')
set_c_type("RANAP Cause", 'uint')
set_c_type("Radio Priority SMS", 'uint')
set_c_type("Radio Priority", 'uint')
set_c_type("Packet Flow Id", 'uint')
set_c_type("Trace Reference", 'uint')
set_c_type("Trace Type", 'uint')
set_c_type("MS Not Reachable Reason", 'uint')
set_c_type("Charging ID", 'uint')
set_c_type("Reordering Required", 'uint')

msg_list["Echo Request"]["table"] = 12
msg_list["Echo Response"]["table"] = 13
#msg_list["Version Not Supported"]["table"] = 3
#msg_list["Node Alive Request"]["table"] = 4
#msg_list["Node Alive Response"]["table"] = 5
#msg_list["Redirection Request"]["table"] = 6
#msg_list["Redirection Response"]["table"] = 7
msg_list["Create PDP Context Request"]["table"] = 15
msg_list["Create PDP Context Response"]["table"] = 16
msg_list["Update PDP Context Request"]["table"] = 17 # FIXME: + 18 " Information Elements in a GGSN-Initiated Update PDP Context Request"
msg_list["Update PDP Context Response"]["table"] = 19 # FIXME: + 20
msg_list["Delete PDP Context Request"]["table"] = 21
msg_list["Delete PDP Context Response"]["table"] = 22
msg_list["Initiate PDP Context Activation Request"]["table"] = 27
msg_list["Initiate PDP Context Activation Response"]["table"] = 28
#msg_list["Error Indication"]["table"] = specified in 3GPP TS 29.281 [42].
msg_list["PDU Notification Request"]["table"] = 23
msg_list["PDU Notification Response"]["table"] = 24
msg_list["PDU Notification Reject Request"]["table"] = 25
msg_list["PDU Notification Reject Response"]["table"] = 26
msg_list["Supported Extension Headers Notification"]["table"] = 14
msg_list["Send Routeing Information for GPRS Request"]["table"] = 29
msg_list["Send Routeing Information for GPRS Response"]["table"] = 30
msg_list["Failure Report Request"]["table"] = 31
msg_list["Failure Report Response"]["table"] = 32
msg_list["Note MS GPRS Present Request"]["table"] = 33
msg_list["Note MS GPRS Present Response"]["table"] = 34
msg_list["Identification Request"]["table"] = 35
msg_list["Identification Response"]["table"] = 36
msg_list["SGSN Context Request"]["table"] = 37
msg_list["SGSN Context Response"]["table"] = 38
msg_list["SGSN Context Acknowledge"]["table"] = 39
msg_list["Forward Relocation Request"]["table"] = 40
msg_list["Forward Relocation Response"]["table"] = 41
msg_list["Forward Relocation Complete"]["table"] = 42
msg_list["Relocation Cancel Request"]["table"] = 43
msg_list["Relocation Cancel Response"]["table"] = 44
msg_list["Forward SRNS Context"]["table"] = 45
msg_list["Forward Relocation Complete Acknowledge"]["table"] = 46
msg_list["Forward SRNS Context Acknowledge"]["table"] = 47
msg_list["UE Registration Query Request"]["table"] = 49
msg_list["UE Registration Query Response"]["table"] = 50
msg_list["RAN Information Relay"]["table"] = 48
msg_list["MBMS Notification Request"]["table"] = 51
msg_list["MBMS Notification Response"]["table"] = 52
msg_list["MBMS Notification Reject Request"]["table"] = 53
msg_list["MBMS Notification Reject Response"]["table"] = 54
msg_list["Create MBMS Context Request"]["table"] = 55
msg_list["Create MBMS Context Response"]["table"] = 56
msg_list["Update MBMS Context Request"]["table"] = 57
msg_list["Update MBMS Context Response"]["table"] = 58
msg_list["Delete MBMS Context Request"]["table"] = 59
msg_list["Delete MBMS Context Response"]["table"] = 60
msg_list["MBMS Registration Request"]["table"] = 61
msg_list["MBMS Registration Response"]["table"] = 62
msg_list["MBMS De-Registration Request"]["table"] = 63
msg_list["MBMS De-Registration Response"]["table"] = 64
msg_list["MBMS Session Start Request"]["table"] = 65
msg_list["MBMS Session Start Response"]["table"] = 66
msg_list["MBMS Session Stop Request"]["table"] = 67
msg_list["MBMS Session Stop Response"]["table"] = 68
msg_list["MBMS Session Update Request"]["table"] = 69
msg_list["MBMS Session Update Response"]["table"] = 70


for key in msg_list.keys():
    if "table" in msg_list[key].keys():
        d_info("[" + key + "]")
        if key == "Delete MBMS Context Request":
            d_info('skipping, broken in source document')
            # FIXME: manually generate the cells for each row
            continue
        cachefile = cachedir + "tlv-msg-" + msg_list[key]["type"] + ".py"
        if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
            exec(open(cachefile).read())
            print("Read from " + cachefile)
        else:
            document = Document(filename)
            f = open(cachefile, 'w')

            ies = []
            write_file(f, "ies = []\n")
            table = document.tables[msg_list[key]["table"]]
            for row in table.rows[1:]:
                cells = get_cells(row.cells)
                if cells is None:
                    continue
                if cells["ie_value"] == "Private Extension":
                    continue;
                ies.append(cells)
                write_cells_to_file("ies", cells)
            msg_list[key]["ies"] = ies
            write_file(f, "msg_list[key][\"ies\"] = ies\n")
            f.close()

f = open(outdir + 'message.h', 'w')
output_header_to_file(f)
f.write("""#if !defined(OGS_GTP_INSIDE) && !defined(OGS_GTP_COMPILATION)
#error "This header cannot be included directly."
#endif

#ifndef OGS_GTP1_MESSAGE_H
#define OGS_GTP1_MESSAGE_H

#ifdef __cplusplus
extern "C" {
#endif

/* 5.1 General format */
#define OGS_GTPV1U_HEADER_LEN   8
#define OGS_GTPV1C_HEADER_LEN   12
#define OGS_GTP1_TEID_LEN        4
typedef struct ogs_gtp1_header_s {
    union {
        struct {
#define OGS_GTP1_VERSION_0 0
#define OGS_GTP1_VERSION_1 1
        ED6(uint8_t version:3;,
            uint8_t pt:1;,
            uint8_t spare1:1;,
            uint8_t e:1;,
            uint8_t s:1;,
            uint8_t pn:1;)
        };
/* GTU-U flags */
#define OGS_GTP1U_FLAGS_V                        0x20
#define OGS_GTP1U_FLAGS_PT                       0x10
#define OGS_GTP1U_FLAGS_E                        0x04
#define OGS_GTP1U_FLAGS_S                        0x02
#define OGS_GTP1U_FLAGS_PN                       0x01
        uint8_t flags;
    };
/* GTP-U message type, defined in 3GPP TS 29.281 Release 11 */
#define OGS_GTP1U_MSGTYPE_ECHO_REQ               1
#define OGS_GTP1U_MSGTYPE_ECHO_RSP               2
#define OGS_GTP1U_MSGTYPE_ERR_IND                26
#define OGS_GTP1U_MSGTYPE_SUPP_EXTHDR_NOTI       31
#define OGS_GTP1U_MSGTYPE_END_MARKER             254
#define OGS_GTP1U_MSGTYPE_GPDU                   255
    uint8_t type;
    uint16_t length;
    uint32_t teid;
    /* Optional based on flags: */
#define OGS_GTP1_XID_TO_SQN(__xid) htobe16((uint16_t)(__xid))
#define OGS_GTP1_SQN_TO_XID(__sqn) be16toh((uint16_t)__sqn)
    uint16_t sqn;
    uint8_t npdu;
    uint8_t ext_hdr_type;
} __attribute__ ((packed)) ogs_gtp1_header_t;

/* GTPv1-C message type */
""")

tmp = [(k, v["type"]) for k, v in msg_list.items()]
sorted_msg_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_msg_list:
    f.write("#define OGS_GTP1_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in type_list.items()]
sorted_type_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_type_list:
    f.write("#define OGS_GTP1_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

f.write("/* Information Element TLV Descriptor */\n")
for (k, v) in sorted_type_list:
    f.write("extern ogs_tlv_desc_t ogs_gtp1_tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Message Descriptor */\n")
for (k, v) in sorted_msg_list:
    f.write("extern ogs_tlv_desc_t ogs_gtp1_tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Structure for Information Element */\n")
for (k, v) in sorted_type_list:
    if  type_list[k].get('c_type', '') == 'uint' and "size" in type_list[k]:
        if type_list[k]["size"] >= 1 and type_list[k]["size"] <= 4:
            f.write("typedef ogs_tlv_uint" + str(type_list[k]["size"]*8) + "_t ogs_gtp1_tlv_" + v_lower(k) + "_t;\n")
        else:
            assert False, "Unknown uint size = %d for key = %s" % (type_list[k]["size"], k)
    else:
        f.write("typedef ogs_tlv_octet_t ogs_gtp1_tlv_" + v_lower(k) + "_t;\n")
f.write("\n")

f.write("/* Structure for Message */\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("typedef struct ogs_gtp1_" + v_lower(k) + "_s {\n")
        for ies in msg_list[k]["ies"]:
            f.write("    ogs_gtp1_tlv_" + v_lower(ie_reference2type(ies["reference"])) + "_t " + \
                    v_lower(ies["ie_value"]) + ";\n")
        f.write("} ogs_gtp1_" + v_lower(k) + "_t;\n")
        f.write("\n")

f.write("typedef struct ogs_gtp1_message_s {\n")
f.write("   ogs_gtp1_header_t h;\n")
f.write("   union {\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        ogs_gtp1_" + v_lower(k) + "_t " + v_lower(k) + ";\n");
f.write("   };\n");
f.write("} ogs_gtp1_message_t;\n\n")

f.write("""int ogs_gtp1_parse_msg(ogs_gtp1_message_t *gtp1_message, ogs_pkbuf_t *pkbuf);
ogs_pkbuf_t *ogs_gtp1_build_msg(ogs_gtp1_message_t *gtp1_message);

#ifdef __cplusplus
}
#endif

#endif /* OGS_GTP1_MESSAGE_H */
""")
f.close()

f = open(outdir + 'message.c', 'w')
output_header_to_file(f)
f.write("""#include "ogs-gtp.h"

""")

for (k, v) in sorted_type_list:
    f.write("ogs_tlv_desc_t ogs_gtp1_tlv_desc_%s =\n" % v_lower(k))
    f.write("{\n")
    if type_list[k]["format"] == 'TV':
        if not "size" in type_list[k]:
            assert False, "Unknown size for TV key = %s" % k
        if type_list[k].get('c_type', '') == 'uint':
            if type_list[k]["size"] >= 1 and type_list[k]["size"] <= 4:
                f.write("    OGS_TV_UINT%d,\n" % (type_list[k]["size"]*8))
            else:
                assert False, "Unsupported %s size = %d for key = %s" % (type_list[k]["c_type"], type_list[k]["size"], k)
        else:
            f.write("    OGS_TV_FIXED_STR,\n")
    elif type_list[k]["format"] == 'TLV':
        if not "size" in type_list[k]:
            f.write("    OGS_TLV_VAR_STR,\n")
        else:
            if type_list[k].get('c_type', '') == 'uint':
                if type_list[k]["size"] >= 1 and type_list[k]["size"] <= 4:
                    f.write("    OGS_TLV_UINT%d,\n" % (type_list[k]["size"]*8))
                else:
                    assert False, "Unsupported %s size = %d for key = %s" % (type_list[k]["c_type"], type_list[k]["size"], k)
            else:
                f.write("    OGS_TLV_FIXED_STR,\n")
    else:
        f.write("    OGS_TLV_VAR_STR,\n")
    f.write("    \"%s\",\n" % k)
    f.write("    OGS_GTP1_%s_TYPE,\n" % v_upper(k))
    if "size" in type_list[k]:
        f.write("    %d,\n" % type_list[k]["size"])
    else:
        f.write("    0,\n")
    f.write("    0,\n")
    f.write("    sizeof(ogs_gtp1_tlv_%s_t),\n" % v_lower(k))
    f.write("    { NULL }\n")
    f.write("};\n\n")

for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("ogs_tlv_desc_t ogs_gtp1_tlv_desc_%s =\n" % v_lower(k))
        f.write("{\n")
        f.write("    OGS_TLV_MESSAGE,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    0, 0, 0, 0, {\n")
        for ies in msg_list[k]["ies"]:
            #if k == 'Create PDP Context Request' and ies["ie_value"] == 'SGSN Address for user traffic':
            #    f.write("        &ogs_tlv_desc_more1,\n")
            #else:
                f.write("        &ogs_gtp1_tlv_desc_%s,\n" % v_lower(ie_reference2type(ies["reference"])))
        f.write("    NULL,\n")
        f.write("}};\n\n")
f.write("\n")

f.write("""int ogs_gtp1_parse_msg(ogs_gtp1_message_t *gtp_message, ogs_pkbuf_t *pkbuf)
{
    int rv = OGS_ERROR;
    ogs_gtp1_header_t *h = NULL;
    uint16_t size = 0;

    ogs_assert(gtp_message);
    ogs_assert(pkbuf);
    ogs_assert(pkbuf->len);

    h = (ogs_gtp1_header_t *)pkbuf->data;
    ogs_assert(h);

    memset(gtp_message, 0, sizeof(ogs_gtp1_message_t));

    if (h->e || h->s || h->pn)
        size = OGS_GTPV1C_HEADER_LEN;
    else
        size = OGS_GTPV1C_HEADER_LEN - 4;

    ogs_assert(ogs_pkbuf_pull(pkbuf, size));
    memcpy(&gtp_message->h, pkbuf->data - size, size);

    gtp_message->h.teid = be32toh(gtp_message->h.teid);

    if (pkbuf->len == 0) {
        ogs_assert(ogs_pkbuf_push(pkbuf, size));
        return OGS_OK;
    }

    switch(gtp_message->h.type) {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("    case OGS_GTP1_%s_TYPE:\n" % v_upper(k))
        f.write("        rv = ogs_tlv_parse_msg_desc(&gtp_message->%s,\n" % v_lower(k))
        f.write("                &ogs_gtp1_tlv_desc_%s, pkbuf, OGS_TLV_MODE_T1_L2);\n" % v_lower(k))
        f.write("        break;\n")
f.write("""    default:
        ogs_warn("Not implmeneted(type:%d)", gtp_message->h.type);
        break;
    }

    ogs_assert(ogs_pkbuf_push(pkbuf, size));

    return rv;
}

""")

f.write("""ogs_pkbuf_t *ogs_gtp1_build_msg(ogs_gtp1_message_t *gtp_message)
{
    ogs_pkbuf_t *pkbuf = NULL;

    ogs_assert(gtp_message);
    switch(gtp_message->h.type) {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("    case OGS_GTP1_%s_TYPE:\n" % v_upper(k))
        f.write("        pkbuf = ogs_tlv_build_msg(&ogs_gtp1_tlv_desc_%s,\n" % v_lower(k))
        f.write("                &gtp_message->%s, OGS_TLV_MODE_T1_L2);\n" % v_lower(k))
        f.write("        break;\n")
f.write("""    default:
        ogs_warn("Not implmeneted(type:%d)", gtp_message->h.type);
        break;
    }

    return pkbuf;
}
""")

f.write("\n")

f.close()
