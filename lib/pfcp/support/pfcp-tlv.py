# Copyright (C) 2019-2023 by Sukchan Lee <acetcom@gmail.com>

# This file is part of Open5GS.

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

from docx import Document
import re, os, sys, string
import datetime
import getopt
import getpass

version = "0.1.0"

msg_list = {}
type_list = {}
group_list = {}

verbosity = 0
filename = ""
outdir = './'
cachedir = './cache/'

FAIL = '\033[91m'
INFO = '\033[93m'
ENDC = '\033[0m'

def d_print(string):
    if verbosity > 0:
        sys.stdout.write(string)

def d_info(string):
    sys.stdout.write(INFO + string + ENDC + "\n")

def d_error(string):
    sys.stderr.write(FAIL + string + ENDC + "\n")
    sys.exit(0)

def write_file(f, string):
    f.write(string)
    d_print(string)

def output_header_to_file(f):
    now = datetime.datetime.now()
    f.write("""/*
 * Copyright (C) 2019-2023 by Sukchan Lee <acetcom@gmail.com>
 *
 * This file is part of Open5GS.
 *
 * This program is free software: you can redistribute it and/or modify
 * it under the terms of the GNU Affero General Public License as published by
 * the Free Software Foundation, either version 3 of the License, or
 * (at your option) any later version.
 *
 * This program is distributed in the hope that it will be useful,
 * but WITHOUT ANY WARRANTY; without even the implied warranty of
 * MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
 * GNU General Public License for more details.
 *
 * You should have received a copy of the GNU General Public License
 * along with this program.  If not, see <https://www.gnu.org/licenses/>.
 */

""")
    f.write("/*******************************************************************************\n")
    f.write(" * This file had been created by pfcp-tlv.py script v%s\n" % (version))
    f.write(" * Please do not modify this file but regenerate it via script.\n")
    f.write(" * Created on: %s by %s\n * from %s\n" % (str(now), getpass.getuser(), filename))
    f.write(" ******************************************************************************/\n\n")

def usage():
    print("Python generating TLV build/parser for PFCP v%s" % (version))
    print("Usage: python pfcp-tlv.py [options]")
    print("Available options:")
    print("-d        Enable script debug")
    print("-f [file] Input file to parse")
    print("-o [dir]  Output files to given directory")
    print("-c [dir]  Cache files to given directory")
    print("-h        Print this help and return")

def v_upper(v):
    return re.sub('5GS', 'FiveGS', re.sub('3GPP', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).upper()))

def v_lower(v):
    return re.sub('5gs', 'fivegs', re.sub('3gpp', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).lower()))

def get_cells(cells):
    note = cells[0].text
    if note.find('NOTE') != -1:
        return None
    comment = cells[2].text.encode('ascii', 'ignore').decode('utf-8')
    comment = re.sub('\n|\"|\'|\\\\', '', comment);
    #print(comment)
    ie_type = re.sub('\s*$', '', re.sub('\'\s*\n*\s*\(NOTE.*\)*', '', cells[-1].text))

    #if ie_type.find('Usage Report') != -1:
    if ie_type == 'Usage Report':
        if comment.find('Report Type') != -1:
            ie_type = "Usage Report Session Report Request"
        elif comment.find('Query URR') != -1:
            ie_type = "Usage Report Session Modification Response"
        elif comment.find('provisioned ') != -1:
            ie_type = "Usage Report Session Deletion Response"
        else:
             assert False, "Unknown IE type : [Usage Report]"

    if ie_type == 'Update BAR':
        if comment.find('7.5.4.11-1') != -1:
            ie_type = "Update BAR Session Modification Request"
        elif comment.find('7.5.9.2-1') != -1:
            ie_type = "Update BAR PFCP Session Report Response"
        else:
             assert False, "Unknown IE type : [Update BAR]"

    if ie_type.find('PFD Contents') != -1:
        ie_type = 'PFD contents'
    elif ie_type.find('UE IP address') != -1:
        ie_type = 'UE IP Address'
    elif ie_type.find('SxSMReq-Flags') != -1:
        ie_type = 'PFCPSMReq-Flags'
    elif ie_type.find('PFCPSRRsp-Flags2') != -1:
        ie_type = 'PFCPSRRsp-Flags'
    elif ie_type.find('IPv4 Configuration Parameters (IP4CP)') != -1:
        ie_type = 'IP4CP'
    elif ie_type.find('QoS Information') != -1:
        ie_type = 'QoS Information in GTP-U Path QoS Report'
    elif ie_type.find('IP Multicast Addressing Info') != -1:
        ie_type = 'IP Multicast Addressing Info within PFCP Session Establishment Request'
    elif ie_type.find('Redundant Transmission Detection Parameters') != -1:
        ie_type = 'Redundant Transmission Parameters'
    elif ie_type.find('MAC Addressed Detected') != -1:
        ie_type = 'MAC Addresses Detected'
    elif ie_type.find('Join IP Multicast Information') != -1:
        ie_type = 'Join IP Multicast Information IE within Usage Report'
    elif ie_type.find('Leave IP Multicast Information') != -1:
        ie_type = 'Leave IP Multicast Information IE within Usage Report'
    elif ie_type.find('PFCP Session Retention Information') != -1:
        ie_type = 'PFCP Session Retention Information within PFCP Association Setup Request'
    elif ie_type.find('GTP-U Path QoS Report') != -1:
        ie_type = 'GTP-U Path QoS Report PFCP Node Report Request'
    elif ie_type.find('Provide RDS configuration information') != -1:
        ie_type = 'Provide RDS Configuration Information'
    elif ie_type.find('RDS configuration information') != -1:
        ie_type = 'RDS Configuration Information'
    elif ie_type.find('Group Id') != -1:
        ie_type = 'Group ID'
    elif ie_type.find('Group-Id') != -1:
        ie_type = 'Group ID'
    elif ie_type.find('L2TP session Indications') != -1:
        ie_type = 'L2TP Session Indications'
    elif ie_type.find('L2TP User Authentication') != -1:
        ie_type = 'L2TP User Authentication IE'
    elif ie_type.find('IP Address and Port Number Replacement') != -1:
        ie_type = 'IP Address and Port number Replacement'
    elif ie_type.find('User Plane Node Management Information Container') != -1:
        ie_type = 'Bridge Management Information Container'
    elif ie_type.find('TSC Management Information') != -1:
        ie_type = 'TSC Management Information IE within PFCP Session Modification Request'
    elif ie_type.find('Query Packet Rate Status') != -1:
        ie_type = 'Query Packet Rate Status IE within PFCP Session Modification Request'
    if ie_type not in type_list.keys():
        assert False, "Unknown IE type : [" \
                + cells[-1].text + "]" + "(" + ie_type + ")"
    presence = cells[1].text
    ie_value = re.sub('\s*\n*\s*\([^\)]*\)*', '', cells[0].text)
    ie_value = re.sub('\n', '', ie_value)
    if ie_value[len(ie_value)-1] == ' ':
        ie_value = ie_value[:len(ie_value)-1]

    tlv_more = "0"  # PFCP has no tlv_more
    if ie_type == 'Create PDR' or ie_type == 'Created PDR' or ie_type == 'Update PDR' or ie_type == "Remove PDR":
        tlv_more = "15"
    if ie_type == 'Create FAR' or ie_type == 'Update FAR' or ie_type == "Remove FAR":
        tlv_more = "15"
    if ie_type == 'Create URR' or ie_type == 'Update URR' or ie_type == "Remove URR":
        tlv_more = "15"
    if ie_type == 'Create QER' or ie_type == 'Update QER' or ie_type == "Remove QER":
        tlv_more = "3"
    if ie_type == 'User Plane IP Resource Information':
        tlv_more = "3"
    if ie_type == 'SDF Filter':
        tlv_more = "7"
    if (ie_type == 'Framed-Route' or
        ie_type == 'Framed-IPv6-Route'):
        tlv_more = "7"
    if (ie_type == 'Usage Report Session Report Request' or
        ie_type == 'Usage Report Session Deletion Response' or
        ie_type == 'Usage Report Session Modification Response'):
        tlv_more = "7"
    if ie_type == 'URR ID' and comment.find('Several IEs within the same IE type may be present') != -1:
        tlv_more = "7"

    if int(tlv_more) > int(type_list[ie_type]["max_tlv_more"]):
        type_list[ie_type]["max_tlv_more"] = tlv_more
        write_file(f, "type_list[\"" + ie_type + "\"][\"max_tlv_more\"] = \"" + tlv_more + "\"\n")

    return { "ie_type" : ie_type, "ie_value" : ie_value, "presence" : presence, "tlv_more" : tlv_more, "comment" : comment }

def write_cells_to_file(name, cells):
    write_file(f, name + ".append({ \"ie_type\" : \"" + cells["ie_type"] + \
        "\", \"ie_value\" : \"" + cells["ie_value"] + \
        "\", \"presence\" : \"" + cells["presence"] + \
        "\", \"tlv_more\" : \"" + cells["tlv_more"] + \
        "\", \"comment\" : \"" + cells["comment"] + "\"})\n")

try:
    opts, args = getopt.getopt(sys.argv[1:], "df:ho:c:", ["debug", "file", "help", "output", "cache"])
except getopt.GetoptError as err:
    # print help information and exit:
    usage()
    sys.exit(2)

for o, a in opts:
    if o in ("-d", "--debug"):
        verbosity = 1
    if o in ("-f", "--file"):
        filename = a
    if o in ("-o", "--output"):
        outdir = a
        if outdir.rfind('/') != len(outdir):
            outdir += '/'
    if o in ("-c", "--cache"):
        cachedir = a
        if cachedir.rfind('/') != len(cachedir):
            cachedir += '/'
    if o in ("-h", "--help"):
        usage()
        sys.exit(2)

if os.path.isfile(filename) and os.access(filename, os.R_OK):
    file = open(filename, 'r')
else:
    d_error("Cannot find file : " + filename)

d_info("[Message List]")
cachefile = cachedir + 'tlv-msg-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    exec(open(cachefile).read())
    print("Read from " + cachefile)
else:
    document = Document(filename)
    f = open(cachefile, 'w')

    msg_table = ""
    for i, table in enumerate(document.tables):
        try:
            cell = table.rows[0].cells[0]
        except:
            continue;
        else:
            if cell.text.find('Message Type value') != -1:
                msg_table = table
                d_print("Table Index = %d\n" % i)

    for row in msg_table.rows[2:-3]:
        key = row.cells[1].text
        type = row.cells[0].text
        if type.isdigit() is False:
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\s*\n*\s*\([^\)]*\)*', '', key)
        key = re.sub('\n', '', key)
        msg_list[key] = { "type": type }
        write_file(f, "msg_list[\"" + key + "\"] = { \"type\" : \"" + type + "\" }\n")
    f.close()

d_info("[IE Type List]")
cachefile = cachedir + 'tlv-type-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    exec(open(cachefile).read())
    print("Read from " + cachefile)
else:
    document = Document(filename)
    f = open(cachefile, 'w')

    ie_table = ""
    for i, table in enumerate(document.tables):
        try:
            cell = table.rows[0].cells[0]
        except:
            pass
        else:
            if cell.text.find('IE Type value') != -1:
                ie_table = table
                d_print("Table Index = %d\n" % i)

    for row in ie_table.rows[1:]:
        key = row.cells[1].text
        type = row.cells[0].text
        if type.isdigit() is False:
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\(', '', key)
        key = re.sub('\)', '', key)
        key = re.sub('\s*$', '', key)

        type_list[key] = { "type": type , "max_tlv_more" : "0" }
        write_file(f, "type_list[\"" + key + "\"] = { \"type\" : \"" + type)
        write_file(f, "\", \"max_tlv_more\" : \"0\" }\n")
    f.close()

d_info("[Group IE List]")
cachefile = cachedir + 'tlv-group-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    exec(open(cachefile).read())
    print("Read from " + cachefile)
else:
    document = Document(filename)
    f = open(cachefile, 'w')

    for i, table in enumerate(document.tables):
        try:
            cell = table.rows[0].cells[0]
        except:
            pass
        else:
            if cell.text.find('Octet') != -1 and \
               table.rows[0].cells[1].text.find('Outer Header to be created') == -1:

                num = 0;
                if len(table.rows[0].cells) > 2 and table.rows[0].cells[2].text.find('IE Type') != -1:
                    num = 2
                elif len(table.rows[0].cells) > 3 and table.rows[0].cells[3].text.find('IE Type') != -1:
                    num = 3
                elif len(table.rows[0].cells) > 4 and table.rows[0].cells[4].text.find('IE Type') != -1:
                    num = 4

                if num == 0:
                    continue;

                row = table.rows[0];

                d_print("Table Index = %d[%d:%s]\n" % (i, num, row.cells[num].text))

                if len(re.findall('\d+', row.cells[num].text)) == 0:
                    continue;
                ie_type = re.findall('\d+', row.cells[num].text)[-1]
                ie_name = re.sub('\s*IE Type.*', '', row.cells[num].text)

                d_print("TYPE:%s NAME:%s\n" % (ie_type, ie_name))

                # SKIP Access Forwarding Action Information
                if (int(ie_type) == 78):
                    ie_name =  "Usage Report Session Modification Response"
                elif (int(ie_type) == 79):
                    ie_name =  "Usage Report Session Deletion Response"
                elif (int(ie_type) == 80):
                    ie_name =  "Usage Report Session Report Request"
                elif (int(ie_type) == 86):
                    ie_name =  "Update BAR Session Modification Request"
                elif (int(ie_type) == 12):
                    ie_name =  "Update BAR PFCP Session Report Response"
                elif (int(ie_type) == 183):
                    ie_name =  "PFCP Session Retention Information within PFCP Association Setup Request"
                elif (int(ie_type) == 188):
                    ie_name =  "IP Multicast Addressing Info within PFCP Session Establishment Request"
                elif (int(ie_type) == 189):
                    ie_name =  "Join IP Multicast Information IE within Usage Report"
                elif (int(ie_type) == 190):
                    ie_name =  "Leave IP Multicast Information IE within Usage Report"
                elif (int(ie_type) == 199):
                    ie_name =  "TSC Management Information IE within PFCP Session Modification Request"
                elif (int(ie_type) == 200):
                    ie_name =  "TSC Management Information IE within PFCP Session Modification Response"
                elif (int(ie_type) == 201):
                    ie_name =  "TSC Management Information IE within PFCP Session Report Request"
                elif (int(ie_type) == 239):
                    ie_name =  "GTP-U Path QoS Report PFCP Node Report Request"
                elif (int(ie_type) == 240):
                    ie_name =  "QoS Information in GTP-U Path QoS Report"
                elif (int(ie_type) == 255):
                    ie_name =  "Redundant Transmission Parameters"
                elif (int(ie_type) == 263):
                    ie_name =  "Query Packet Rate Status IE within PFCP Session Modification Request"
                elif (int(ie_type) == 264):
                    ie_name =  "Packet Rate Status Report IE within PFCP Session Modification Response"

                if ie_name.find('Non-3GPP Access Forwarding Action Information') != -1:
                    ie_idx = str(int(ie_type)+100)
                    group_list[ie_name] = { "index" : ie_idx, "type" : ie_type, "ies" : ies }
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"index\" : \"" + ie_idx + "\", \"type\" : \"" + ie_type + "\", \"ies\" : ies }\n")
                    continue

                if ie_name not in group_list.keys():
                    ies = []
                    write_file(f, "ies = []\n")
                    for row in table.rows[4:]:
                        cells = get_cells(row.cells)
                        if cells is None:
                            continue

                        ies.append(cells)
                        write_cells_to_file("ies", cells)

                    ie_idx = str(int(ie_type)+100)
                    group_list[ie_name] = { "index" : ie_idx, "type" : ie_type, "ies" : ies }
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"index\" : \"" + ie_idx + "\", \"type\" : \"" + ie_type + "\", \"ies\" : ies }\n")
    f.close()

msg_list["PFCP Heartbeat Request"]["table"] = 9
msg_list["PFCP Heartbeat Response"]["table"] = 10
msg_list["PFCP PFD Management Request"]["table"] = 11
msg_list["PFCP PFD Management Response"]["table"] = 14
msg_list["PFCP Association Setup Request"]["table"] = 15
msg_list["PFCP Association Setup Response"]["table"] = 20
msg_list["PFCP Association Update Request"]["table"] = 21
msg_list["PFCP Association Update Response"]["table"] = 23
msg_list["PFCP Association Release Request"]["table"] = 24
msg_list["PFCP Association Release Response"]["table"] = 25
msg_list["PFCP Version Not Supported Response"]["table"] = 0
msg_list["PFCP Node Report Request"]["table"] = 26
msg_list["PFCP Node Report Response"]["table"] = 33
msg_list["PFCP Session Set Deletion Request"]["table"] = 34
msg_list["PFCP Session Set Deletion Response"]["table"] = 35
msg_list["PFCP Session Set Modification Request"]["table"] = 36
msg_list["PFCP Session Set Modification Response"]["table"] = 38
msg_list["PFCP Session Establishment Request"]["table"] = 39
msg_list["PFCP Session Establishment Response"]["table"] = 72
msg_list["PFCP Session Modification Request"]["table"] = 86
msg_list["PFCP Session Modification Response"]["table"] = 112
msg_list["PFCP Session Deletion Request"]["table"] = 117
msg_list["PFCP Session Deletion Response"]["table"] = 118
msg_list["PFCP Session Report Request"]["table"] = 121
msg_list["PFCP Session Report Response"]["table"] = 133

for key in msg_list.keys():
    if "table" in msg_list[key].keys():
        d_info("[" + key + "]")
        cachefile = cachedir + "tlv-msg-" + msg_list[key]["type"] + ".py"
        if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
            exec(open(cachefile).read())
            print("Read from " + cachefile)
        else:
            document = Document(filename)
            f = open(cachefile, 'w')

            table = document.tables[msg_list[key]["table"]]
#            if key.find('Association') != -1:
#                start_i = 1
#            elif key.find('Heartbeat') != -1:
#                start_i = 1
#            else:
#                start_i = 2
            start_i = 2

            ies = []
            write_file(f, "ies = []\n")
            if key != "PFCP Session Deletion Request" and key != "PFCP Version Not Supported Response":
                for row in table.rows[start_i:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    item = {
                        "ie_type" : cells["ie_type"],
                        "ie_value" : cells["ie_value"],
                        "presence" : cells["presence"],
                        "tlv_more" : cells["tlv_more"],
                        "comment" : cells["comment"]
                        }
                    ies.append(item)
                    write_cells_to_file("ies", item)

            msg_list[key]["ies"] = ies
            write_file(f, "msg_list[key][\"ies\"] = ies\n")
            f.close()

type_list["Cause"]["size"] = 1                              # Type 19
type_list["Source Interface"]["size"] = 1                   # Type 20
type_list["Gate Status"]["size"] = 1                        # Type 25
type_list["QER Correlation ID"]["size"] = 4                 # Type 28
type_list["Precedence"]["size"] = 4                         # Type 29
type_list["Time Threshold"]["size"] = 4                     # Type 32
type_list["Reporting Triggers"]["size"] = 3                 # Type 37
type_list["Report Type"]["size"] = 1                        # Type 39
type_list["Offending IE"]["size"] = 2                       # Type 40
type_list["Destination Interface"]["size"] = 1              # Type 42
type_list["Apply Action"]["size"] = 2                       # Type 44
type_list["PFCPSMReq-Flags"]["size"] = 1                    # Type 49
type_list["PFCPSRRsp-Flags"]["size"] = 1                    # Type 50
type_list["PDR ID"]["size"] = 2                             # Type 56
type_list["Measurement Method"]["size"] = 1                 # Type 62
type_list["Usage Report Trigger"]["size"] = 3               # Type 63
type_list["Measurement Period"]["size"] = 4                 # Type 64
type_list["Duration Measurement"]["size"] = 4               # Type 67
type_list["Time of First Packet"]["size"] = 4               # Type 69
type_list["Time of Last Packet"]["size"] = 4                # Type 70
type_list["Quota Holding Time"]["size"] = 4                 # Type 71
type_list["Time Quota"]["size"] = 4                         # Type 74
type_list["Start Time"]["size"] = 4                         # Type 75
type_list["End Time"]["size"] = 4                           # Type 76
type_list["URR ID"]["size"] = 4                             # Type 81
type_list["BAR ID"]["size"] = 1                             # Type 88
type_list["CP Function Features"]["size"] = 1               # Type 89
type_list["Recovery Time Stamp"]["size"] = 4                # Type 96
type_list["UR-SEQN"]["size"] = 4                            # Type 104
type_list["FAR ID"]["size"] = 4                             # Type 108
type_list["QER ID"]["size"] = 4                             # Type 109
type_list["PDN Type"]["size"] = 1                           # Type 113
type_list["RQI"]["size"] = 1                                # Type 123
type_list["QFI"]["size"] = 1                                # Type 124
type_list["Event Quota"]["size"] = 4                        # Type 148
type_list["Event Threshold"]["size"] = 4                    # Type 149
type_list["Averaging Window"]["size"] = 4                   # Type 157
type_list["Paging Policy Indicator"]["size"] = 1            # Type 158
type_list["3GPP Interface Type"]["size"] = 1                # Type 160
type_list["PFCPSRReq-Flags"]["size"] = 1                    # Type 161
type_list["PFCPAUReq-Flags"]["size"] = 1                    # Type 162
type_list["Quota Validity Time"]["size"] = 4                # Type 181
type_list["PFCPSEReq-Flags"]["size"] = 1                    # Type 186
type_list["Data Status"]["size"] = 1                        # Type 260

f = open(outdir + 'message.h', 'w')
output_header_to_file(f)
f.write("""#if !defined(OGS_PFCP_INSIDE) && !defined(OGS_PFCP_COMPILATION)
#error "This header cannot be included directly."
#endif

#ifndef OGS_PFCP_MESSAGE_H
#define OGS_PFCP_MESSAGE_H

#ifdef __cplusplus
extern "C" {
#endif

/* 5.1 General format */
#define OGS_PFCP_HEADER_LEN 16
#define OGS_PFCP_SEID_LEN   8
typedef struct ogs_pfcp_header_s {
    union {
        struct {
        ED4(uint8_t version:3;,
            uint8_t spare1:3;,
            uint8_t mp:1;,
            uint8_t seid_presence:1;)
        };
        uint8_t flags;
    };
    uint8_t type;
    uint16_t length;
    union {
        struct {
            uint64_t seid;
            /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
#define OGS_PFCP_XID_TO_SQN(__xid) htobe32(((__xid) << 8))
#define OGS_PFCP_SQN_TO_XID(__sqn) (be32toh(__sqn) >> 8)
            uint32_t sqn;
        };
        /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
        uint32_t sqn_only;
    };
} __attribute__ ((packed)) ogs_pfcp_header_t;

/* PFCP message type */
""")

tmp = [(k, v["type"]) for k, v in msg_list.items()]
sorted_msg_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_msg_list:
    f.write("#define OGS_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in type_list.items()]
sorted_type_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_type_list:
    f.write("#define OGS_PFCP_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

f.write("/* Information Element TLV Descriptor */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    f.write("extern ogs_tlv_desc_t ogs_pfcp_tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

for k, v in group_list.items():
    if v_lower(k) == "ethernet_packet_filter":
        v["index"] = "1"
    if v_lower(k) == "redundant_transmission_parameters":
        v["index"] = "2"
    if v_lower(k) == "ip_multicast_addressing_info_within_pfcp_session_establishment_request":
        v["index"] = "3"
    if v_lower(k) == "pdi":
        v["index"] = "4"
    if v_lower(k) == "transport_delay_reporting":
        v["index"] = "5"
    if v_lower(k) == "create_pdr":
        v["index"] = "6"
    if v_lower(k) == "forwarding_parameters":
        v["index"] = "7"
    if v_lower(k) == "duplicating_parameters":
        v["index"] = "8"
    if v_lower(k) == "redundant_transmission_forwarding_parameters":
        v["index"] = "9"
    if v_lower(k) == "mbs_multicast_parameters":
        v["index"] = "10"
    if v_lower(k) == "add_mbs_unicast_parameters":
        v["index"] = "11"
    if v_lower(k) == "create_far":
        v["index"] = "12"
    if v_lower(k) == "update_forwarding_parameters":
        v["index"] = "13"
    if v_lower(k) == "update_duplicating_parameters":
        v["index"] = "14"
    if v_lower(k) == "update_far":
        v["index"] = "15"
    if v_lower(k) == "pfd_context":
        v["index"] = "16"
    if v_lower(k) == "application_id_s_pfds":
        v["index"] = "17"
    if v_lower(k) == "ethernet_traffic_information":
        v["index"] = "18"
    if v_lower(k) == "_access_forwarding_action_information":
        v["index"] = "19"
    if v_lower(k) == "non__access_forwarding_action_information":
        v["index"] = "20"
    if v_lower(k) == "update__access_forwarding_action_information":
        v["index"] = "21"
    if v_lower(k) == "update_non__access_forwarding_action_information":
        v["index"] = "22"
    if v_lower(k) == "access_availability_report":
        v["index"] = "23"
    if v_lower(k) == "qos_monitoring_report":
        v["index"] = "24"
    if v_lower(k) == "mptcp_parameters":
        v["index"] = "25"
    if v_lower(k) == "atsss_ll_parameters":
        v["index"] = "26"
    if v_lower(k) == "pmf_parameters":
        v["index"] = "27"
    if v_lower(k) == "join_ip_multicast_information_ie_within_usage_report":
        v["index"] = "28"
    if v_lower(k) == "leave_ip_multicast_information_ie_within_usage_report":
        v["index"] = "29"

tmp = [(k, v["index"]) for k, v in group_list.items()]
sorted_group_list = sorted(tmp, key=lambda tup: int(tup[1]), reverse=False)

f.write("/* Group Information Element TLV Descriptor */\n")
for (k, v) in sorted_group_list:
    f.write("extern ogs_tlv_desc_t ogs_pfcp_tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Message Descriptor */\n")
for (k, v) in sorted_msg_list:
    f.write("extern ogs_tlv_desc_t ogs_pfcp_msg_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Structure for Information Element */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    if "size" in type_list[k]:
        if type_list[k]["size"] == 1:
            f.write("typedef ogs_tlv_uint8_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 2:
            f.write("typedef ogs_tlv_uint16_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 3:
            f.write("typedef ogs_tlv_uint24_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 4:
            f.write("typedef ogs_tlv_uint32_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        else:
            assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
    else:
        f.write("typedef ogs_tlv_octet_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
f.write("\n")

tmp = []
f.write("/* Structure for Group Information Element */\n")
for (k, v) in sorted_group_list:
    f.write("typedef struct ogs_pfcp_tlv_" + v_lower(k) + "_s {\n")
    f.write("    ogs_tlv_presence_t presence;\n")
    for ies in group_list[k]["ies"]:
        if type_list[ies["ie_type"]]["max_tlv_more"] != "0" and ies["tlv_more"] != "0":
            f.write("    ogs_pfcp_tlv_" + v_lower(ies["ie_type"]) + "_t " + v_lower(ies["ie_value"]) + "[" + str(int(ies["tlv_more"])+1) + "];\n")
        else:
            f.write("    ogs_pfcp_tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                    v_lower(ies["ie_value"]) + ";\n")
    f.write("} ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
    f.write("\n")

f.write("/* Structure for Message */\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("typedef struct ogs_" + v_lower(k) + "_s {\n")
        for ies in msg_list[k]["ies"]:
            if type_list[ies["ie_type"]]["max_tlv_more"] != "0" and ies["tlv_more"] != "0":
                f.write("    ogs_pfcp_tlv_" + v_lower(ies["ie_type"]) + "_t " + v_lower(ies["ie_value"]) + "[" + str(int(ies["tlv_more"])+1) + "];\n")
            else:
                f.write("    ogs_pfcp_tlv_" + v_lower(ies["ie_type"]) + "_t " + v_lower(ies["ie_value"]) + ";\n")

        f.write("} ogs_" + v_lower(k) + "_t;\n")
        f.write("\n")

f.write("typedef struct ogs_pfcp_message_s {\n")
f.write("   ogs_pfcp_header_t h;\n")
f.write("   union {\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        ogs_" + v_lower(k) + "_t " + v_lower(k) + ";\n");
f.write("   };\n");
f.write("} ogs_pfcp_message_t;\n\n")

f.write("""ogs_pfcp_message_t *ogs_pfcp_parse_msg(ogs_pkbuf_t *pkbuf);
void ogs_pfcp_message_free(ogs_pfcp_message_t *pfcp_message);
ogs_pkbuf_t *ogs_pfcp_build_msg(ogs_pfcp_message_t *pfcp_message);

#ifdef __cplusplus
}
#endif

#endif /* OGS_PFCP_MESSAGE_H */
""")
f.close()

f = open(outdir + 'message.c', 'w')
output_header_to_file(f)
f.write("""#include "ogs-pfcp.h"

""")

for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    f.write("ogs_tlv_desc_t ogs_pfcp_tlv_desc_%s =\n" % v_lower(k))
    f.write("{\n")
    if "size" in type_list[k]:
        if type_list[k]["size"] == 1:
            f.write("    OGS_TLV_UINT8,\n")
        elif type_list[k]["size"] == 2:
            f.write("    OGS_TLV_UINT16,\n")
        elif type_list[k]["size"] == 3:
            f.write("    OGS_TLV_UINT24,\n")
        elif type_list[k]["size"] == 4:
            f.write("    OGS_TLV_UINT32,\n")
        else:
            assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
    else:
        f.write("    OGS_TLV_VAR_STR,\n")
    f.write("    \"%s\",\n" % k)
    f.write("    OGS_PFCP_%s_TYPE,\n" % v_upper(k))
    if "size" in type_list[k]:
        f.write("    %d,\n" % type_list[k]["size"])
    else:
        f.write("    0,\n")
    f.write("    0,\n")
    f.write("    sizeof(ogs_pfcp_tlv_%s_t),\n" % v_lower(k))
    f.write("    { NULL }\n")
    f.write("};\n\n")

for (k, v) in sorted_group_list:
    f.write("ogs_tlv_desc_t ogs_pfcp_tlv_desc_%s =\n" % v_lower(k))
    f.write("{\n")
    f.write("    OGS_TLV_COMPOUND,\n")
    f.write("    \"%s\",\n" % k)
    f.write("    OGS_PFCP_%s_TYPE,\n" % v_upper(k))
    f.write("    0,\n")
    f.write("    0,\n")
    f.write("    sizeof(ogs_pfcp_tlv_%s_t),\n" % v_lower(k))
    f.write("    {\n")
    for ies in group_list[k]["ies"]:
        f.write("        &ogs_pfcp_tlv_desc_%s,\n" % v_lower(ies["ie_type"]))
        if type_list[ies["ie_type"]]["max_tlv_more"] != "0" and ies["tlv_more"] != "0":
            f.write("        &ogs_tlv_desc_more" + str(int(ies["tlv_more"])+1) + ",\n")
    f.write("        NULL,\n")
    f.write("    }\n")
    f.write("};\n\n")

for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("ogs_tlv_desc_t ogs_pfcp_msg_desc_%s =\n" % v_lower(k))
        f.write("{\n")
        f.write("    OGS_TLV_MESSAGE,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    0, 0, 0, 0, {\n")
        for ies in msg_list[k]["ies"]:
            f.write("        &ogs_pfcp_tlv_desc_%s,\n" % v_lower(ies["ie_type"]))
            if type_list[ies["ie_type"]]["max_tlv_more"] != "0" and ies["tlv_more"] != "0":
                f.write("        &ogs_tlv_desc_more" + str(int(ies["tlv_more"])+1) + ",\n")
        f.write("    NULL,\n")
        f.write("}};\n\n")
f.write("\n")

f.write("""ogs_pfcp_message_t *ogs_pfcp_parse_msg(ogs_pkbuf_t *pkbuf)
{
    int rv = OGS_ERROR;
    ogs_pfcp_header_t *h = NULL;
    uint16_t size = 0;

    ogs_pfcp_message_t *pfcp_message = NULL;

    ogs_assert(pkbuf);
    ogs_assert(pkbuf->len);

    h = (ogs_pfcp_header_t *)pkbuf->data;
    ogs_assert(h);

    pfcp_message = ogs_calloc(1, sizeof(*pfcp_message));
    if (!pfcp_message) {
        ogs_error("No memory");
        return NULL;
    }

    if (h->seid_presence)
        size = OGS_PFCP_HEADER_LEN;
    else
        size = OGS_PFCP_HEADER_LEN-OGS_PFCP_SEID_LEN;

    if (ogs_pkbuf_pull(pkbuf, size) == NULL) {
        ogs_error("ogs_pkbuf_pull() failed [len:%d]", pkbuf->len);
        ogs_pfcp_message_free(pfcp_message);
        return NULL;
    }
    memcpy(&pfcp_message->h, pkbuf->data - size, size);

    if (h->seid_presence) {
        pfcp_message->h.seid = be64toh(pfcp_message->h.seid);
    } else {
        pfcp_message->h.sqn = pfcp_message->h.sqn_only;
    }

    if (pkbuf->len == 0)
        return pfcp_message;

    switch(pfcp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case OGS_%s_TYPE:\n" % v_upper(k))
        if k != "PFCP Session Deletion Request" and k != "PFCP Version Not Supported Response":
            f.write("            rv = ogs_tlv_parse_msg(&pfcp_message->%s,\n" % v_lower(k))
            f.write("                    &ogs_pfcp_msg_desc_%s, pkbuf, OGS_TLV_MODE_T2_L2);\n" % v_lower(k))
            f.write("            ogs_expect(rv == OGS_OK);\n")
        f.write("            break;\n")
f.write("""        default:
            ogs_warn("Not implemented(type:%d)", pfcp_message->h.type);
            break;
    }

    if (rv != OGS_OK) {
        ogs_pfcp_message_free(pfcp_message);
        return NULL;
    }

    return pfcp_message;
}

void ogs_pfcp_message_free(ogs_pfcp_message_t *pfcp_message)
{
    ogs_assert(pfcp_message);
    ogs_free(pfcp_message);
}

""")

f.write("""ogs_pkbuf_t *ogs_pfcp_build_msg(ogs_pfcp_message_t *pfcp_message)
{
    ogs_pkbuf_t *pkbuf = NULL;

    ogs_assert(pfcp_message);
    switch(pfcp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case OGS_%s_TYPE:\n" % v_upper(k))
        f.write("            pkbuf = ogs_tlv_build_msg(&ogs_pfcp_msg_desc_%s,\n" % v_lower(k))
        f.write("                    &pfcp_message->%s, OGS_TLV_MODE_T2_L2);\n" % v_lower(k))
        f.write("            break;\n")
f.write("""        default:
            ogs_warn("Not implemented(type:%d)", pfcp_message->h.type);
            break;
    }

    return pkbuf;
}
""")

f.write("\n")

f.close()
